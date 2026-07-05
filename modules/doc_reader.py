from docx import Document


def read_docx(file_path):

    print(f"正在读取Word：{file_path.name}")

    doc = Document(file_path)

    text = []

    # 段落
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)

    # 表格
    for table in doc.tables:
        for row in table.rows:

            row_data = []

            for cell in row.cells:

                value = cell.text.strip()

                if value:
                    row_data.append(value)

            if row_data:
                text.append(" | ".join(row_data))

    result = "\n".join(text)

    print(f"读取完成，共 {len(result)} 个字符")

    return result